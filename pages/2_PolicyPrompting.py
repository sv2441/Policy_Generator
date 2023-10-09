import streamlit as st
import docx2txt
import pandas as pd
import base64
import csv
import math
import docx
import os
from langchain.output_parsers import OutputFixingParser
from langchain.llms import OpenAI
from langchain.chains import LLMChain, SimpleSequentialChain
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.chat_models import ChatOpenAI
from langchain.output_parsers import ResponseSchema, StructuredOutputParser
from dotenv import load_dotenv


load_dotenv()
os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
# os.environ["OPENAI_API_KEY"] = os.getenv('OPENAI_API_KEY')

chat_llm = ChatOpenAI(temperature=0.0)

st.title("Policy Prompting")

@st.cache_resource
def generation(prompt,target):
    if os.path.exists('Policy.doc'):
        doc = docx.Document('Policy.doc')
    else:
        doc = docx.Document()
    title_template = prompt
    prompt = ChatPromptTemplate.from_template(template=title_template)
    messages = prompt.format_messages(topic=target)
    response = chat_llm(messages)
    content = str(response.content)
    doc.add_paragraph(content)
    doc.save('Policy.doc')
    result=convert_doc_to_text('Policy.doc')
    st.write(result)
    st.cache_resource.clear()
    
    with open('Policy.doc', 'rb') as f:
        doc_data = f.read()
    b64 = base64.b64encode(doc_data).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64}" download="policy.doc">Download Result</a>'
    st.markdown(href, unsafe_allow_html=True)
    
 

# Function to convert DOC to Text
@st.cache_resource
def convert_doc_to_text(doc_file):
    text = docx2txt.process(doc_file)
    return text

target=" "

# Text area for user input
st.subheader("User Prompt")
prompt = st.text_area("Enter your Prompt here:")

# Save user input to a file
if st.button("Submit"):
    generation(prompt,target)
    st.cache_resource.clear()
    



